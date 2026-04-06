import io
try:
    from docx import Document
except ImportError:
    Document = None

def generate_text_report(stats):
    """
    Generates a raw .txt report.
    """
    lines = [
        "SENTENCE PROFILER REPORT",
        "========================\n",
        f"Overall Assessment: {stats['overall_complexity']}\n",
        "1. Sentence Length Analysis:",
        f"   - Total Words: {stats['length_stats']['total_words']}",
        f"   - Total Sentences: {stats['length_stats']['total_sentences']}",
        f"   - Average Sentence Length: {stats['length_stats']['avg_length']} words\n",
        "2. Grammatical Structure:",
        f"   - Total Clauses: {stats['grammar_stats']['total_clauses']}",
        f"   - Subordinate Clauses: {stats['grammar_stats']['subordinate_clauses']}\n",
        "3. Noun Phrase (NP) Complexity:",
        f"   - Total Noun Phrases: {stats['np_stats']['total_nps']}",
        f"   - Average NP Length: {stats['np_stats']['avg_np_length']} words",
        f"   - Total Adjective Modifiers: {stats['np_stats']['total_modifiers']}\n",
        "4. Linguistic Features:",
        f"   - Unique Base Words: {stats['feature_stats']['unique_words']}",
        f"   - Lexical Diversity: {stats['feature_stats']['lexical_diversity']} (Scale 0-1)"
    ]
    return "\n".join(lines)

def generate_word_report(stats):
    """
    Generates a .docx report and returns a Byte Stream. 
    """
    if not Document:
        raise Exception("python-docx is not installed.")
        
    doc = Document()
    doc.add_heading('Sentence Profiler Analysis Report', 0)
    
    doc.add_heading('Overall Result', level=1)
    doc.add_paragraph(f"Complexity Classification: {stats['overall_complexity']}")
    
    doc.add_heading('1. Sentence Length Analysis', level=2)
    doc.add_paragraph(f"Total Words: {stats['length_stats']['total_words']}")
    doc.add_paragraph(f"Total Sentences: {stats['length_stats']['total_sentences']}")
    doc.add_paragraph(f"Average Sentence Length: {stats['length_stats']['avg_length']} words")
    
    doc.add_heading('2. Grammatical Structure', level=2)
    doc.add_paragraph(f"Total Clauses: {stats['grammar_stats']['total_clauses']}")
    doc.add_paragraph(f"Subordinate Clauses: {stats['grammar_stats']['subordinate_clauses']}")
    
    doc.add_heading('3. Noun Phrase (NP) Complexity', level=2)
    doc.add_paragraph(f"Total Noun Phrases: {stats['np_stats']['total_nps']}")
    doc.add_paragraph(f"Average NP Length: {stats['np_stats']['avg_np_length']} words")
    doc.add_paragraph(f"Total Adjective Modifiers: {stats['np_stats']['total_modifiers']}")
    
    doc.add_heading('4. Linguistic Features', level=2)
    doc.add_paragraph(f"Unique Base Words: {stats['feature_stats']['unique_words']}")
    doc.add_paragraph(f"Lexical Diversity Score: {stats['feature_stats']['lexical_diversity']} (Scale 0-1)")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
